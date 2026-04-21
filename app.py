from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, send_from_directory, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from oauthlib.oauth2 import WebApplicationClient
import requests
import os
import json
from datetime import datetime
from docx import Document
from docx.shared import Inches
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors


app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-change-in-production'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///resumecraft.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Google OAuth configuration
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID', 'your-google-client-id')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET', 'your-google-client-secret')
GOOGLE_DISCOVERY_URL = "https://accounts.google.com/.well-known/openid_configuration"

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'signin'

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# User Model
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    name = db.Column(db.String(120), nullable=False)
    google_id = db.Column(db.String(120), unique=True)
    password_hash = db.Column(db.String(128))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Resume Model
class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    data = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# OAuth client setup
client = WebApplicationClient(GOOGLE_CLIENT_ID)

def get_google_provider_cfg():
    return requests.get(GOOGLE_DISCOVERY_URL).json()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/builder')
@login_required
def builder():
    return render_template('builder.html')

@app.route('/signin')
def signin():
    if current_user.is_authenticated:
        return redirect(url_for('builder'))
    return render_template('signin.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

# Authentication routes
@app.route('/login', methods=['POST'])
def login():
    email = request.form.get('email')
    password = request.form.get('password')

    if not email or not password:
        return jsonify({'error': 'Email and password are required'}), 400

    user = User.query.filter_by(email=email).first()

    if user and user.password_hash and check_password_hash(user.password_hash, password):
        login_user(user)
        return jsonify({'success': True})
    else:
        return jsonify({'error': 'Invalid email or password'}), 401

@app.route('/signup', methods=['POST'])
def signup():
    name = request.form.get('name')
    email = request.form.get('email')
    password = request.form.get('password')
    confirm_password = request.form.get('confirm_password')

    # Validation
    if not name or not email or not password or not confirm_password:
        return jsonify({'error': 'All fields are required'}), 400

    if password != confirm_password:
        return jsonify({'error': 'Passwords do not match'}), 400

    if len(password) < 6:
        return jsonify({'error': 'Password must be at least 6 characters long'}), 400

    existing_user = User.query.filter_by(email=email).first()
    if existing_user:
        return jsonify({'error': 'Email already registered'}), 400

    # Create new user with hashed password
    hashed_password = generate_password_hash(password)
    user = User(
        email=email,
        name=name,
        password_hash=hashed_password
    )
    
    try:
        db.session.add(user)
        db.session.commit()
        login_user(user)
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': 'Failed to create account. Please try again.'}), 500

# Google OAuth routes
@app.route('/auth/google')
def google_login():
    google_provider_cfg = get_google_provider_cfg()
    authorization_endpoint = google_provider_cfg["authorization_endpoint"]

    request_uri = client.prepare_request_uri(
        authorization_endpoint,
        redirect_uri=request.base_url + "/callback",
        scope=["openid", "email", "profile"],
    )
    return redirect(request_uri)

@app.route('/auth/google/callback')
def google_callback():
    code = request.args.get("code")

    google_provider_cfg = get_google_provider_cfg()
    token_endpoint = google_provider_cfg["token_endpoint"]

    token_url, headers, body = client.prepare_token_request(
        token_endpoint,
        authorization_response=request.url,
        redirect_url=request.base_url,
        code=code
    )

    token_response = requests.post(
        token_url,
        headers=headers,
        data=body,
        auth=(GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET),
    )

    client.parse_request_body_response(json.dumps(token_response.json()))

    userinfo_endpoint = google_provider_cfg["userinfo_endpoint"]
    uri, headers, body = client.add_token(userinfo_endpoint)
    userinfo_response = requests.get(uri, headers=headers, data=body)

    if userinfo_response.json().get("email_verified"):
        google_id = userinfo_response.json()["sub"]
        email = userinfo_response.json()["email"]
        name = userinfo_response.json()["name"]

        user = User.query.filter_by(google_id=google_id).first()
        if not user:
            user = User.query.filter_by(email=email).first()
            if user:
                user.google_id = google_id
            else:
                user = User(google_id=google_id, email=email, name=name)
                db.session.add(user)
            db.session.commit()

        login_user(user)
        return redirect(url_for('builder'))

    return "User email not available or not verified by Google.", 400

# API routes
@app.route('/api/user')
def api_user():
    if current_user.is_authenticated:
        return jsonify({'user': {
            'id': current_user.id,
            'email': current_user.email,
            'name': current_user.name
        }})
    return jsonify({'user': None})

@app.route('/api/resume', methods=['POST'])
@login_required
def save_resume():
    data = request.get_json()
    resume = Resume(user_id=current_user.id, data=json.dumps(data))
    db.session.add(resume)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/resumes')
@login_required
def get_resumes():
    resumes = Resume.query.filter_by(user_id=current_user.id).order_by(Resume.created_at.desc()).all()
    return jsonify({'resumes': [{
        'id': r.id,
        'data': json.loads(r.data),
        'created_at': r.created_at.isoformat()
    } for r in resumes]})

@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    resume_data = request.get_json()

    doc = Document()
    doc.add_heading(resume_data.get('fullName', 'Full Name'), 0)

    # Personal Information
    doc.add_heading('Personal Information', level=1)
    doc.add_paragraph(f"Job Title: {resume_data.get('jobTitle', '')}")
    doc.add_paragraph(f"Email: {resume_data.get('email', '')}")
    doc.add_paragraph(f"Phone: {resume_data.get('phone', '')}")
    doc.add_paragraph(f"Location: {resume_data.get('location', '')}")

    # Professional Summary
    if resume_data.get('summary'):
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(resume_data['summary'])

    # Work Experience - Updated for array
    if resume_data.get('experiences') and len(resume_data['experiences']) > 0:
        doc.add_heading('Work Experience', level=1)
        for exp in resume_data['experiences']:
            experience = f"{exp.get('role', '')} at {exp.get('company', '')}"
            start = exp.get('start', '')
            end = exp.get('end', 'Present')
            if start or end:
                experience += f" ({start} - {end})"
            doc.add_heading(experience, level=2)
            if exp.get('details'):
                doc.add_paragraph(exp['details'])
            if exp.get('projects'):
                doc.add_paragraph(f"Projects: {exp['projects']}")

    # Education
    if resume_data.get('educations') and len(resume_data['educations']) > 0:
        doc.add_heading('Education', level=1)
        for edu in resume_data['educations']:
            education = f"{edu.get('degree', '')}, {edu.get('college', '')}"
            year = edu.get('year', '')
            grade = edu.get('grade', '')
            if year or grade:
                education += f" ({year}, {grade})"
            doc.add_paragraph(education)

    # Skills
    all_skills = []
    if resume_data.get('skills', {}).get('tech'):
        all_skills.extend(resume_data['skills']['tech'])
    if resume_data.get('skills', {}).get('soft'):
        all_skills.extend(resume_data['skills']['soft'])
    if all_skills:
        doc.add_heading('Skills', level=1)
        for skill in all_skills:
            doc.add_paragraph(skill.strip(), style='List Bullet')

    # Certifications
    if resume_data.get('certifications') or resume_data.get('achievements'):
        doc.add_heading('Certifications & Achievements', level=1)
        if resume_data.get('certifications'):
            doc.add_paragraph(resume_data['certifications'])
        if resume_data.get('achievements'):
            doc.add_paragraph(resume_data['achievements'])

    # Save to BytesIO
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    return send_file(
        bio,
        as_attachment=True,
        download_name='resume.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/generate-pdf', methods=['POST'])
def generate_pdf():
    data = request.get_json()
    resume_data = data.get('resumeData', {})
    style_settings = data.get('styleSettings', {})

    # Create PDF with ReportLab
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1,  # Center
        textColor=colors.darkblue
    )
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=12,
        textColor=colors.darkblue
    )

    story = []

    # Title
    story.append(Paragraph(resume_data.get('fullName', 'Full Name'), title_style))
    story.append(Spacer(1, 12))

    # Subtitle
    if resume_data.get('jobTitle'):
        story.append(Paragraph(resume_data.get('jobTitle'), styles['Heading3']))
        story.append(Spacer(1, 12))

    # Contact info
    contact = []
    if resume_data.get('email'):
        contact.append(resume_data['email'])
    if resume_data.get('phone'):
        contact.append(resume_data['phone'])
    if resume_data.get('location'):
        contact.append(resume_data['location'])
    if contact:
        story.append(Paragraph(" | ".join(contact), styles['Normal']))
        story.append(Spacer(1, 24))

    # Summary
    if resume_data.get('summary'):
        story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        story.append(Paragraph(resume_data['summary'], styles['Normal']))
        story.append(Spacer(1, 18))

    # Experience
    if resume_data.get('experiences'):
        story.append(Paragraph("EXPERIENCE", heading_style))
        for exp in resume_data['experiences']:
            role_company = f"<b>{exp.get('role', '')}</b><br/><i>{exp.get('company', '')}</i>"
            dates = f"{exp.get('start', '')} - {exp.get('end', 'Present')}"
            story.append(Paragraph(role_company, styles['Normal']))
            story.append(Paragraph(dates, styles['Italic']))
            if exp.get('details'):
                story.append(Paragraph(exp['details'], styles['Normal']))
            story.append(Spacer(1, 6))
        story.append(Spacer(1, 18))

    # Education
    if resume_data.get('educations'):
        story.append(Paragraph("EDUCATION", heading_style))
        for edu in resume_data['educations']:
            edu_info = f"<b>{edu.get('degree', '')}</b><br/>{edu.get('college', '')}"
            details = f"({edu.get('year', '')}, {edu.get('grade', '')})"
            story.append(Paragraph(edu_info, styles['Normal']))
            story.append(Paragraph(details, styles['Italic']))
            story.append(Spacer(1, 6))
        story.append(Spacer(1, 18))

    # Skills
    skills_data = resume_data.get('skills', {})
    all_skills = skills_data.get('tech', []) + skills_data.get('soft', [])
    if all_skills:
        story.append(Paragraph("SKILLS", heading_style))
        skills_text = ", ".join(all_skills)
        story.append(Paragraph(skills_text, styles['Normal']))
        story.append(Spacer(1, 18))

    # Languages
    if resume_data.get('languages'):
        story.append(Paragraph("LANGUAGES", heading_style))
        story.append(Paragraph(resume_data['languages'], styles['Normal']))
        story.append(Spacer(1, 18))

    # Certifications
    if resume_data.get('certifications') or resume_data.get('achievements'):
        story.append(Paragraph("CERTIFICATIONS & ACHIEVEMENTS", heading_style))
        if resume_data.get('certifications'):
            story.append(Paragraph(resume_data['certifications'], styles['Normal']))
        if resume_data.get('achievements'):
            story.append(Paragraph(resume_data['achievements'], styles['Normal']))

    doc.build(story)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='resume.pdf',
        mimetype='application/pdf'
    )

@app.route('/api/upload-resume', methods=['POST'])
def upload_resume():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['resume']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file:
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        return jsonify({
            'filename': filename,
            'path': filepath
        })

    return jsonify({'error': 'File upload failed'}), 400

# Static file serving for assets
@app.route('/assets/<path:filename>')
def assets(filename):
    return app.send_static_file(f'assets/{filename}')

@app.route('/data/<path:filename>')
def data_files(filename):
    return send_from_directory(os.path.join(app.root_path, 'data'), filename)

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True, port=3000)